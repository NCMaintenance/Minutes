import streamlit as st
import google.generativeai as genai
import json
import os
import time
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import urllib.request
import io
import tempfile
import re
import struct
import pandas as pd
import altair as alt
from google.api_core.exceptions import ResourceExhausted, ServiceUnavailable, PermissionDenied

# --- Configuration ---
GEMINI_MODEL_NAME = 'gemini-3-flash-preview'
TTS_MODEL_NAME = 'gemini-2.5-flash-preview-tts'
LOGO_URL = "https://www.esther.ie/wp-content/uploads/2022/05/HSE-Logo-Green-NEW-no-background.png"
FAVICON_URL = "https://assets.hse.ie/static/hse-frontend/assets/favicons/favicon.ico"

# --- API Key Management ---
def get_available_keys():
    keys = []
    key_names = ["GEMINI_API_KEY", "GEMINI_API_KEY2", "GEMINI_API_KEY3"]
    for name in key_names:
        if name in st.secrets:
            keys.append(st.secrets[name])
    if not keys:
        st.error("No API Keys found in secrets. Please add GEMINI_API_KEY.")
        st.stop()
    return keys

if "key_index" not in st.session_state:
    st.session_state.key_index = 0

def configure_genai_with_current_key():
    keys = get_available_keys()
    if st.session_state.key_index >= len(keys):
        st.session_state.key_index = 0
    genai.configure(api_key=keys[st.session_state.key_index])
    return genai.GenerativeModel(model_name=GEMINI_MODEL_NAME)

# --- Helper: Safe Response Extractor ---
def safe_get_text(response):
    try:
        if not response.candidates: return None
        candidate = response.candidates[0]
        if candidate.content.parts: return candidate.content.parts[0].text
        return None
    except Exception: return None

# --- Helper: Detect Speakers (Cached) ---
@st.cache_data
def detect_speakers(text):
    """Finds speaker labels like '**Speaker 1**:' or 'Speaker 1:'"""
    if not text: return []
    # Regex for bolded or plain speaker labels at start of lines
    pattern = r'(?m)^(?:[\*\_]{2})?([A-Za-z0-9\s\(\)\-\.]+?)(?:[\*\_]{2})?[:]'
    matches = re.findall(pattern, text)
    return sorted(list(set(matches)))

# --- Helper: Add WAV Header ---
def add_wav_header(pcm_data, sample_rate=24000, channels=1, bit_depth=16):
    header = b'RIFF'
    header += struct.pack('<I', 36 + len(pcm_data))
    header += b'WAVEfmt '
    header += struct.pack('<I', 16)
    header += struct.pack('<H', 1)
    header += struct.pack('<H', channels)
    header += struct.pack('<I', sample_rate)
    header += struct.pack('<I', sample_rate * channels * (bit_depth // 8))
    header += struct.pack('<H', channels * (bit_depth // 8))
    header += struct.pack('<H', bit_depth)
    header += b'data'
    header += struct.pack('<I', len(pcm_data))
    return header + pcm_data

# --- Robust Audio Processor ---
def process_audio_with_rotation(tmp_file_path, context_info):
    max_retries = 6 
    base_delay = 1
    keys = get_available_keys()
    
    context_str = f"Context: {context_info}" if context_info else ""
    prompt = f"""
    You are a precise transcription engine for the Health Service Executive (HSE) Ireland.
    {context_str}
    Task: Output the raw transcription of this audio. 
    Constraint: Do NOT include preamble like "Here is the transcript". Do NOT include markdown blocks. Just the dialogue.
    Language: Strict Irish English spelling (e.g. 'Programme', 'Paediatric', 'Centre', 'Realise', 'Colour').
    Format:
    **Speaker Name**: Text...
    **Speaker Name**: Text...
    """

    for attempt in range(max_retries):
        audio_file = None
        try:
            model = configure_genai_with_current_key()
            if attempt > 0: st.toast(f"Retry {attempt}...", icon="🔄")
            audio_file = genai.upload_file(path=tmp_file_path, display_name="HSE_Audio")
            
            while audio_file.state.name == "PROCESSING":
                time.sleep(2)
                audio_file = genai.get_file(audio_file.name)
            
            if audio_file.state.name == "FAILED": raise Exception("Audio processing failed.")

            response = model.generate_content([prompt, audio_file], request_options={"timeout": 1200})
            text = safe_get_text(response)
            
            try: genai.delete_file(audio_file.name)
            except: pass
            
            # FIX 6: Guard against empty or failed partial transcripts
            if text and len(text.strip()) > 20: 
                return text
            elif text:
                 raise Exception("Response too short (potential error)")
            else: 
                 raise Exception("Empty response from AI")

        except Exception as e:
            if audio_file:
                try: genai.delete_file(audio_file.name)
                except: pass
            
            st.session_state.key_index = (st.session_state.key_index + 1) % len(keys)
            time.sleep(base_delay * (1.5 ** attempt))
            
    raise Exception("System busy. Please try again.")

# --- Robust Text Generator ---
def robust_text_gen(prompt):
    max_retries = 6
    keys = get_available_keys()
    
    for attempt in range(max_retries):
        try:
            model = configure_genai_with_current_key()
            response = model.generate_content(prompt, request_options={"timeout": 600})
            text = safe_get_text(response)
            if text: return text
        except Exception:
            pass
        
        st.session_state.key_index = (st.session_state.key_index + 1) % len(keys)
        time.sleep(1)
        
    raise Exception("Unable to generate text.")

# --- Audio Generator (Podcast) ---
def generate_podcast_audio(script_text):
    configure_genai_with_current_key()
    try:
        model = genai.GenerativeModel(model_name=TTS_MODEL_NAME)
        prompt = f"Read this naturally:\n{script_text}"
        response = model.generate_content(
            prompt,
            generation_config={
                "response_modalities": ["AUDIO"],
                "speech_config": {
                    "voice_config": {"prebuilt_voice_config": {"voice_name": "Aoede"}}
                }
            }
        )
        if response.candidates and response.candidates[0].content.parts:
            for part in response.candidates[0].content.parts:
                if part.inline_data:
                    return part.inline_data.data, part.inline_data.mime_type
        return None, None
    except Exception as e:
        st.warning(f"Audio Unavailable: {e}")
        return None, None

# --- Minutes Structure ---
def generate_hse_minutes(structured):
    now = datetime.now()
    def get(val, default="Not stated"): return val if val and str(val).strip().lower() != "not mentioned" else default
    def bullets(val):
        if isinstance(val, list) and val:
            items = [item for item in val if str(item).strip() and str(item).strip().lower() != "not mentioned"]
            if items: return "".join([f"• {item}\n" for item in items])
        return "• None recorded\n"

    # Added extra newlines before Signature block
    template = f"""HSE Capital & Estates Meeting Minutes
Meeting Title: {get(structured.get("meetingTitle"), "Meeting")}
Date: {get(structured.get("meetingDate"), now.strftime("%d/%m/%Y"))}
Time: {get(structured.get("startTime"), "00:00")} - {get(structured.get("endTime"), "00:00")}
Location: {get(structured.get("location"))}
Chairperson: {get(structured.get("chairperson"))}
Minute Taker: {get(structured.get("minuteTaker"))}
________________________________________
1. Attendance
Present:
{bullets(structured.get("attendees", []))}
Apologies:
{bullets(structured.get("apologies", []))}
________________________________________
2. Minutes of Previous Meeting / Matters Arising
{bullets(structured.get("mattersArising", []))}
________________________________________
3. Declarations of Interest
• {get(structured.get("declarationsOfInterest"), "None declared.")}
________________________________________
4. Capital Projects Update
4.1 Major Projects (Capital)
{bullets(structured.get("majorProjects", []))}
4.2 Minor Works / Equipment / ICT
{bullets(structured.get("minorProjects", []))}
________________________________________
5. Estates Strategy and Planning
{bullets(structured.get("estatesStrategy", []))}
________________________________________
6. Health & Safety / Regulatory Compliance
{bullets(structured.get("healthSafety", []))}
________________________________________
7. Risk Register
{bullets(structured.get("riskRegister", []))}
________________________________________
8. Finance Update
{bullets(structured.get("financeUpdate", []))}
________________________________________
9. AOB
{bullets(structured.get("aob", []))}
________________________________________
10. Next Meeting
• {get(structured.get("nextMeetingDate"))}
________________________________________



Minutes Approved By: ____________________ Date: ___________
"""
    return template

def create_docx(content, kind="minutes"):
    doc = Document()
    
    # 1. Add HSE Logo
    try:
        # Download logo to temp file
        req = urllib.request.Request(LOGO_URL, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as response:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_logo:
                tmp_logo.write(response.read())
                tmp_logo_path = tmp_logo.name
        
        doc.add_picture(tmp_logo_path, width=Inches(1.2))
        os.remove(tmp_logo_path)
    except Exception:
        pass # Fallback if no internet or url fail

    # 2. Define Styles with HSE Green
    styles = doc.styles
    HSE_GREEN = RGBColor(0, 86, 59)
    
    # Update Heading 1 style
    h1 = styles['Heading 1']
    h1.font.color.rgb = HSE_GREEN
    h1.font.size = Pt(16)
    h1.font.bold = True
    
    # Update Heading 2 style
    h2 = styles['Heading 2']
    h2.font.color.rgb = HSE_GREEN
    h2.font.size = Pt(13)
    h2.font.bold = True

    # 3. Parse content lines for smart formatting
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            # Add small spacing for empty lines, but not too much
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            continue
            
        # Skip visual separators in the DOCX (we use style/headers instead)
        if "________" in line and "Approved By" not in line:
            continue 
        
        # Detect Main Title
        if "HSE Capital & Estates Meeting Minutes" in line:
            doc.add_heading(line, level=1)
        
        # Detect Section Headers (e.g., "1. Attendance")
        elif re.match(r'^\d+\.\s', line):
            doc.add_heading(line, level=2)
            
        # Detect Sub-headers (e.g., "4.1 Major Projects")
        elif re.match(r'^\d+\.\d+\s', line):
             p = doc.add_paragraph()
             runner = p.add_run(line)
             runner.bold = True
             runner.font.color.rgb = HSE_GREEN
        
        # Detect Key-Value pairs (Date: ..., Time: ...) for bolding
        elif ":" in line and len(line.split(":")[0]) < 40 and not line.startswith("•"):
            parts = line.split(":", 1)
            p = doc.add_paragraph()
            p.add_run(parts[0] + ":").bold = True
            p.add_run(parts[1])
            
        # Signature Block specific formatting
        elif "Minutes Approved By:" in line:
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(36) # Extra space before signature
            p.add_run(line).bold = True
            
        else:
            doc.add_paragraph(line)
            
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- Setup ---
st.set_page_config(page_title="HSE MAI Recap", layout="wide", page_icon=FAVICON_URL)

# FIX 1 (Styling): CSS to make Radio buttons look like tabs (PREMIUM GLASS LOOK)
st.markdown("""
<style>
    /* Global Typography & Colors */
    h1, h2, h3, h4 { color: #00563B !important; }
    
    /* Standard Streamlit Button Override */
    .stButton > button { 
        background-color: #00563B !important; 
        color: white !important; 
        border: none !important;
        border-radius: 8px !important;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        transition: transform 0.2s, box-shadow 0.2s;
    }
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 12px rgba(0,0,0,0.15);
        background-color: #007a53 !important;
    }
    
    /* Sidebar Background */
    div[data-testid="stSidebar"] { background-color: #f8f9fa; }
    .stInfo { background-color: #e8f5e9; color: #00563B; }
    
    /* --- Premium Glass Tab-Like Radio Buttons --- */
    
    /* Hide the default radio circle/dot */
    div[role="radiogroup"] > label > div:first-child {
        display: none;
    }
    
    /* Container styling for horizontal alignment */
    div[role="radiogroup"] {
        background: rgba(255, 255, 255, 0.5);
        display: flex;
        flex-direction: row;
        gap: 8px; /* Tighter gap for tab feel */
        padding: 6px;
        border-radius: 12px;
        overflow-x: auto;
        border: 1px solid rgba(0,0,0,0.05);
    }
    
    /* Individual Tab Styling */
    div[role="radiogroup"] label {
        background: transparent;
        border: 1px solid transparent;
        padding: 8px 20px;
        border-radius: 8px; /* Slightly squarer for tab feel */
        cursor: pointer;
        transition: all 0.2s ease-in-out;
        color: #555;
        font-weight: 500;
        min-width: 100px;
        display: flex;
        align-items: center;
        justify-content: center;
    }
    
    /* Hover State */
    div[role="radiogroup"] label:hover {
        background: rgba(0, 86, 59, 0.05);
        color: #00563B;
    }
    
    /* Selected State - Matching Heading Green */
    div[role="radiogroup"] label[data-checked="true"] {
        background-color: #00563B !important; /* Exact match to Heading */
        color: white !important;
        box-shadow: 0 2px 6px rgba(0, 86, 59, 0.25) !important;
        font-weight: 600;
        border-radius: 8px;
    }
    
    /* --- Metric Card Styling --- */
    .metric-card {
        background-color: white;
        border: 1px solid rgba(0, 86, 59, 0.1);
        border-radius: 12px;
        padding: 20px;
        text-align: center;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        transition: transform 0.2s;
    }
    .metric-card:hover {
        transform: translateY(-3px);
        box-shadow: 0 8px 15px rgba(0,0,0,0.1);
    }
    .metric-value {
        font-size: 32px;
        font-weight: 700;
        color: #00563B;
        margin: 0;
    }
    .metric-label {
        font-size: 14px;
        color: #666;
        margin: 0;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
</style>
""", unsafe_allow_html=True)

try:
    if "GEMINI_API_KEY" in st.secrets:
        configure_genai_with_current_key()
    else:
        st.error("Secrets missing.")
        st.stop()
except:
    pass

if "password_verified" not in st.session_state:
    st.session_state.password_verified = False

# FIX 1: Initialize Persistent Active View state
if "current_view" not in st.session_state:
    st.session_state.current_view = "📄 Transcript"

if not st.session_state.password_verified:
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        st.image(LOGO_URL, width=150)
        st.markdown("### HSE Secure Login")
        with st.form("password_form"):
            user_password = st.text_input("Enter Access Code:", type="password")
            if st.form_submit_button("Login"):
                if st.secrets.get("password") and user_password == st.secrets["password"]:
                    st.session_state.password_verified = True
                    st.rerun()
                elif not st.secrets.get("password"):
                      st.warning("Password not set.")
                else:
                    st.error("Invalid code.")
    st.stop()

if "messages" not in st.session_state: st.session_state.messages = []
if "transcript" not in st.session_state: st.session_state.transcript = ""

# FIX: Initialize detected speakers cache
if "detected_speakers" not in st.session_state:
    st.session_state.detected_speakers = []

# --- Sidebar ---
with st.sidebar:
    st.image(LOGO_URL, width="stretch")
    st.title("MAI Recap")
    
    # 1. Reset
    if st.button("🔄 New Meeting / Reset"):
        # FIX 3: Exclude active view from reset to prevent jumping
        preserve_keys = ['password_verified', 'key_index', 'current_view'] 
        for key in list(st.session_state.keys()):
            if key not in preserve_keys: del st.session_state[key]
        st.rerun()
    
    st.markdown("---")
    
    # 2. Context Input
    st.markdown("### ℹ️ Meeting Context")
    context_info = st.text_area("Details (Chair, Topics):", placeholder="e.g. Chair: Sarah. Topic: Budget.", height=80)
    
    # 3. Speaker Renaming
    if st.session_state.transcript:
        st.markdown("---")
        st.markdown("### 👥 Speaker ID Manager")
        
        # Use cached detected speakers
        detected = st.session_state.detected_speakers
        
        if detected:
            with st.form("speaker_update_form"):
                replacements = {}
                for spk in detected:
                    new_name = st.text_input(f"Rename '{spk}':", placeholder=spk)
                    if new_name and new_name != spk:
                        replacements[spk] = new_name
                
                if st.form_submit_button("Update Transcript"):
                    txt = st.session_state.transcript
                    
                    # FIX 4: Robust Regex Renaming with whitespace handling
                    for old, new in replacements.items():
                        # Replace bold speaker label
                        txt = re.sub(
                            rf"\*\*{re.escape(old)}\*\*",
                            f"**{new}**",
                            txt
                        )
                        # Replace plain speaker label with possible whitespace
                        txt = re.sub(
                            rf"(?m)^\s*{re.escape(old)}:",
                            f"{new}:",
                            txt
                        )
                    
                    st.session_state.transcript = txt
                    
                    # Update cache and sync display key
                    st.session_state.detected_speakers = detect_speakers(txt)
                    st.session_state.transcript_display = txt
                    
                    st.toast("Transcript updated with new names!", icon="✅")
                    st.rerun()
        else:
            st.caption("No speakers detected yet.")

    st.markdown("---")
    if st.button("Created by Dave Maher"):
        st.info("Property of Dave Maher.")
    st.markdown("**Version:** 5.7.0 (Sentiment & Fixes)")

# --- Header ---
c1, c2 = st.columns([1, 6])
with c1: st.image(LOGO_URL, width=120)
with c2: 
    st.title("Meeting Minutes Generator")
    st.markdown("#### Automated Documentation System")

# --- Input ---
# FIX: Changed labels to Nouns to avoid "Action Button" confusion
mode = st.radio("Input Source:", ["Microphone", "File Upload"], horizontal=True)
audio_bytes = None

if mode == "File Upload":
    audio_bytes = st.file_uploader("Upload (WAV, MP3, M4A)", type=["wav", "mp3", "m4a", "ogg"])
    if audio_bytes: st.audio(audio_bytes)
else:
    audio_bytes = st.audio_input("🎙️ Click to Record")
    if audio_bytes: st.audio(audio_bytes)

if audio_bytes and st.button("🧠 Transcribe"):
    with st.spinner("Processing..."):
        if hasattr(audio_bytes, "read"): data = audio_bytes.read()
        else: data = audio_bytes
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        
        try:
            transcript_text = process_audio_with_rotation(tmp_path, context_info)
            
            # Atomic update
            st.session_state["transcript"] = transcript_text
            st.session_state.detected_speakers = detect_speakers(transcript_text)
            st.session_state.transcript_display = transcript_text
            
            # FIX 1 (UX): Switch to transcript view automatically
            st.session_state.current_view = "📄 Transcript"
            
            st.success("Transcription Complete.")
            st.rerun()
        except Exception as e:
            st.error(f"Error: {e}")
        finally:
            if os.path.exists(tmp_path): os.remove(tmp_path)

# --- Output Views ---
if st.session_state.transcript:
    st.markdown("---")
    
    # FIX 1: Radio Button Navigation (Persistent) replacement for st.tabs
    nav_options = ["📄 Transcript", "🏥 Minutes", "📝 Briefing", "🎙️ Podcast", "📊 Analytics", "💬 Chat"]
    
    # Ensure current_view is valid
    if st.session_state.current_view not in nav_options:
        st.session_state.current_view = nav_options[0]
        
    selected_view = st.radio(
        "Navigation", 
        nav_options, 
        key="current_view", 
        horizontal=True,
        label_visibility="collapsed"
    )

    # 1. Transcript
    if selected_view == "📄 Transcript":
        # FIX 2: Sync Manual Edits
        # Initialize key if needed
        if "transcript_display" not in st.session_state:
            st.session_state.transcript_display = st.session_state.transcript
            
        edited_transcript = st.text_area(
            "Full Transcript (Editable):", 
            key="transcript_display", # Binds to st.session_state.transcript_display
            height=500
        )
        
        # Check for divergence (Manual Edit detection)
        if edited_transcript != st.session_state.transcript:
            st.session_state.transcript = edited_transcript
            st.session_state.detected_speakers = detect_speakers(edited_transcript)
            # No rerun needed here, it syncs for next action

    # 2. Minutes
    elif selected_view == "🏥 Minutes":
        if st.button("Generate Minutes", key="btn_min"):
            with st.spinner("Extracting..."):
                prompt = f"""
                Extract structured data from transcript (JSON). 
                Language: Strict Irish English (e.g. 'Paediatric', 'Programme'). Currency: Euro.
                Transcript: {st.session_state.transcript}
                Keys: meetingTitle, meetingDate, startTime, endTime, location, chairperson, minuteTaker, attendees, apologies, mattersArising, declarationsOfInterest, majorProjects, minorProjects, estatesStrategy, healthSafety, riskRegister, financeUpdate, aob, nextMeetingDate.
                """
                try:
                    res = robust_text_gen(prompt)
                    # FIX 5: Safer JSON extraction with fallback
                    try:
                        structured = json.loads(res)
                    except:
                        json_match = re.search(r"({[\s\S]*})", res, re.DOTALL)
                        if json_match:
                            structured = json.loads(json_match.group(1))
                        else:
                            raise Exception("No JSON found in response")
                            
                    st.session_state.minutes = generate_hse_minutes(structured)
                except Exception as e: st.error(f"Error: {e}")
        
        if "minutes" in st.session_state:
            st.text_area("Draft:", st.session_state.minutes, height=600)
            st.download_button("Download DOCX", create_docx(st.session_state.minutes), "Minutes.docx")

    # 3. Briefing
    elif selected_view == "📝 Briefing":
        if st.button("Generate Briefing", key="btn_brief"):
            with st.spinner("Analyzing..."):
                prompt = f"""
                Write a neutral, matter-of-fact Executive Briefing based on this transcript.
                Language: Strict Irish English spelling (e.g. 'Realise', 'Centre', 'Colour').
                Do NOT use corporate fluff. Be candid and objective.
                Sections: Executive Summary, Key Decisions, Critical Risks, Action Items.
                Transcript: {st.session_state.transcript}
                """
                st.session_state.briefing = robust_text_gen(prompt)
        
        if "briefing" in st.session_state:
            st.markdown(st.session_state.briefing)
            st.download_button("Download Briefing", create_docx(st.session_state.briefing), "Briefing.docx")

    # 4. Podcast
    elif selected_view == "🎙️ Podcast":
        st.info("NotebookLM Style: Two neutral analysts discussing the meeting.")
        if st.button("Generate Script", key="btn_script"):
            with st.spinner("Writing script..."):
                prompt = f"""
                Convert this transcript into a podcast script between two hosts (Host and Expert).
                Language: Irish English spelling and phrasing.
                Tone: Candid, neutral, analytical (Like NotebookLM) but with Irish nuances. NOT corporate/PR.
                They should discuss the meeting outcomes naturally, pointing out interesting dynamics or risks.
                Transcript: {st.session_state.transcript}
                """
                st.session_state.podcast = robust_text_gen(prompt)
        
        if "podcast" in st.session_state:
            st.text_area("Script:", st.session_state.podcast, height=300)
            if st.button("Generate Audio", key="btn_audio"):
                with st.spinner("Synthesizing..."):
                    audio, mime = generate_podcast_audio(st.session_state.podcast)
                    if audio:
                        if "pcm" in mime.lower() or "raw" in mime.lower():
                             audio = add_wav_header(audio)
                             mime = "audio/wav"
                        st.session_state.pod_audio = audio
                        st.session_state.pod_mime = mime
                    else: st.error("Audio generation failed.")
        
        if "pod_audio" in st.session_state:
            st.audio(st.session_state.pod_audio, format=st.session_state.pod_mime)

    # 5. Analytics (MOVED HERE)
    elif selected_view == "📊 Analytics":
        st.markdown("### Meeting Analytics")
        
        # Parse transcript for analysis
        txt = st.session_state.transcript
        chunks = re.split(r'(?m)^(?:[\*\_]{2})?([A-Za-z0-9\s\(\)\-\.]+?)(?:[\*\_]{2})?[:]', txt)
        
        if len(chunks) > 1:
            data = []
            total_words = 0
            
            for i in range(1, len(chunks), 2):
                if i+1 < len(chunks):
                    speaker = chunks[i].strip()
                    content = chunks[i+1].strip()
                    word_count = len(content.split())
                    total_words += word_count
                    data.append({"Speaker": speaker, "Words": word_count, "Segment": i//2})
            
            df = pd.DataFrame(data)
            
            # --- Metrics Row ---
            col1, col2, col3 = st.columns(3)
            
            est_minutes = round(total_words / 130)
            if est_minutes < 1: est_minutes = "< 1"
            unique_speakers = df['Speaker'].nunique() if not df.empty else 0
            
            def metric_card(label, value):
                return f"""
                <div class="metric-card">
                    <p class="metric-label">{label}</p>
                    <p class="metric-value">{value}</p>
                </div>
                """
            
            with col1: st.markdown(metric_card("Est. Duration", f"{est_minutes} min"), unsafe_allow_html=True)
            with col2: st.markdown(metric_card("Total Words", f"{total_words}"), unsafe_allow_html=True)
            with col3: st.markdown(metric_card("Active Speakers", f"{unique_speakers}"), unsafe_allow_html=True)
            
            st.markdown("---")
            
            # --- Charts Row 1 ---
            c1, c2 = st.columns([1, 1])
            
            with c1:
                st.markdown("#### Share of Voice")
                if not df.empty:
                    speaker_stats = df.groupby("Speaker")["Words"].sum().reset_index()
                    base = alt.Chart(speaker_stats).encode(
                        theta=alt.Theta("Words", stack=True),
                        color=alt.Color("Speaker", scale=alt.Scale(scheme='greens'))
                    )
                    pie = base.mark_arc(outerRadius=120, innerRadius=60)
                    text = base.mark_text(radius=140).encode(
                        text="Speaker",
                        order=alt.Order("Words", sort="descending")
                    )
                    st.altair_chart(pie + text, width="stretch")
            
            with c2:
                st.markdown("#### Conversation Flow")
                if not df.empty:
                    scatter = alt.Chart(df).mark_circle(size=100).encode(
                        x=alt.X('Segment', title='Timeline (Sequencing)'),
                        y=alt.Y('Speaker', title=None),
                        color=alt.Color('Speaker', legend=None, scale=alt.Scale(scheme='greens')),
                        tooltip=['Speaker', 'Words', 'Segment']
                    ).interactive()
                    st.altair_chart(scatter, width="stretch")
                    
            st.markdown("---")

            # --- Charts Row 2 ---
            c3, c4 = st.columns([1, 1])

            with c3:
                st.markdown("#### Verbosity (Avg Words/Turn)")
                if not df.empty:
                    verbosity = df.groupby("Speaker")["Words"].mean().reset_index()
                    bar = alt.Chart(verbosity).mark_bar().encode(
                        x=alt.X('Words', title='Avg Words per Turn'),
                        y=alt.Y('Speaker', sort='-x'),
                        color=alt.Color('Speaker', legend=None, scale=alt.Scale(scheme='greens')),
                        tooltip=['Speaker', 'Words']
                    )
                    st.altair_chart(bar, width="stretch")

            with c4:
                # Rename to "Meeting Activity" to be accurate
                st.markdown("#### Meeting Activity (Word Volume)")
                if not df.empty:
                    area = alt.Chart(df).mark_area(opacity=0.6, interpolate='step').encode(
                        x=alt.X('Segment', title='Timeline'),
                        y=alt.Y('Words', title='Volume'),
                        color=alt.value('#00563B'),
                        tooltip=['Segment', 'Words', 'Speaker']
                    )
                    st.altair_chart(area, width="stretch")
            
            st.markdown("---")
            
            # --- New Feature: Sentiment Analysis ---
            if st.button("📉 Analyze Tone/Sentiment"):
                with st.spinner("Analyzing emotional arc... (This may take a moment)"):
                    try:
                        sentiment_prompt = f"""
                        Analyze the sentiment of this transcript over the course of the meeting. 
                        Divide the meeting into 10 sequential segments. 
                        For each segment return a JSON object with:
                        - 'Segment': int (1-10)
                        - 'Sentiment': float (-1.0 to 1.0, where -1 is negative/tense, 0 is neutral, 1 is positive)
                        - 'Label': str (e.g. 'Tense', 'Optimistic', 'Neutral', 'Action-Oriented')
                        
                        Return ONLY a JSON list of these objects.
                        Transcript: {st.session_state.transcript[:30000]} 
                        """
                        # Limit transcript length to avoid huge context usage for just sentiment
                        
                        response = robust_text_gen(sentiment_prompt)
                        json_match = re.search(r"(\[[\s\S]*\])", response, re.DOTALL)
                        
                        if json_match:
                            sentiment_data = json.loads(json_match.group(1))
                            st.session_state.sentiment_df = pd.DataFrame(sentiment_data)
                        else:
                            st.error("Could not parse sentiment data.")
                            
                    except Exception as e:
                        st.error(f"Sentiment Analysis Failed: {e}")

            if "sentiment_df" in st.session_state:
                st.markdown("#### 🎭 Emotional Arc (Tone)")
                
                # Color scale condition
                domain = [-1, 0, 1]
                range_ = ['#d32f2f', '#fbc02d', '#388e3c'] # Red, Yellow, Green

                sentiment_chart = alt.Chart(st.session_state.sentiment_df).mark_line(point=True).encode(
                    x=alt.X('Segment', title='Timeline (10 Segments)'),
                    y=alt.Y('Sentiment', title='Sentiment Score (-1 to 1)', scale=alt.Scale(domain=[-1, 1])),
                    color=alt.value('#00563B'),
                    tooltip=['Segment', 'Sentiment', 'Label']
                ).properties(height=300)
                
                # Add a zero line
                rule = alt.Chart(pd.DataFrame({'y': [0]})).mark_rule(color='gray', strokeDash=[5, 5]).encode(y='y')
                
                st.altair_chart(sentiment_chart + rule, width="stretch")

        else:
            st.info("Insufficient data to generate analytics. Please transcribe a meeting first.")

    # 6. Chat
    elif selected_view == "💬 Chat":
        # FIX 7: Chat History Limit (Max 20)
        MAX_CHAT_HISTORY = 20
        if len(st.session_state.messages) > MAX_CHAT_HISTORY:
            st.session_state.messages = st.session_state.messages[-MAX_CHAT_HISTORY:]

        for m in st.session_state.messages:
            with st.chat_message(m["role"]): st.markdown(m["content"])
        
        if q := st.chat_input("Question?"):
            st.session_state.messages.append({"role": "user", "content": q})
            with st.chat_message("user"): st.markdown(q)
            with st.chat_message("assistant"):
                prompt = f"Answer neutrally using Irish English spelling/grammar. Transcript: {st.session_state.transcript}\nQ: {q}"
                ans = robust_text_gen(prompt)
                st.markdown(ans)
                st.session_state.messages.append({"role": "assistant", "content": ans})
# --- Footer ---
st.markdown("---")
st.markdown(
    "**Disclaimer:** This implementation has been tested using sample data. "
    "Adjustments may be required to ensure optimal performance and accuracy with real-world meeting audio. "
    "Always verify the accuracy of transcriptions and minutes."
)
st.markdown("Created by Dave Maher | For HSE internal use.")

